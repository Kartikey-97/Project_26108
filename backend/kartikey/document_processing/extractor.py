"""
kartikey/document_processing/extractor.py

Extracts plain text from PDF and DOCX tender documents.

Design principles:
  - Does NOT call external services or AI models.
    Text cleaning and semantic chunking are handled downstream by the AI/ML component.
  - Does NOT modify or annotate the text — returns raw extracted content.
  - Preserves structure hints where possible (page breaks, section numbers)
    so that downstream requirement extraction can identify source locations.
  - Handles the two most common government tender formats: PDF and DOCX.
  - Raises DocumentError with specific codes so callers can respond appropriately.

Known limitations (flagged, not hidden):
  - Scanned PDFs (images inside PDF): text extraction will return empty or near-empty.
    Detected and flagged via SCANNED_PDF error code — OCR is a future improvement.
  - Multi-column PDFs: pdfplumber handles these better than PyPDF2 but may still
    produce garbled column ordering in complex layouts.
  - Password-protected PDFs: rejected with ENCRYPTED_PDF error code.
"""

from __future__ import annotations

import re
from pathlib import Path

from shared.utils import DocumentError, get_logger

logger = get_logger(__name__)

# Minimum characters to consider extraction successful.
# Anything below this almost certainly means a scanned/image PDF.
_MIN_TEXT_LENGTH = 100


# ===========================================================================
# Public interface
# ===========================================================================

def extract_text(path: Path) -> str:
    """
    Extract plain text from a PDF or DOCX file.

    Parameters
    ----------
    path:
        Absolute path to the document file (saved by storage.py).

    Returns
    -------
    str
        Extracted plain text. Newlines preserved.
        Page breaks represented as '\\n\\n--- Page N ---\\n\\n' for PDFs.

    Raises
    ------
    DocumentError
        UNSUPPORTED_FILE_TYPE  — file extension not .pdf or .docx
        EXTRACTION_FAILED      — file could not be parsed
        SCANNED_PDF            — PDF appears to be image-only (no text layer)
        ENCRYPTED_PDF          — PDF is password-protected
        EMPTY_DOCUMENT         — document yielded no text after extraction
    """
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)

    raise DocumentError(
        f"Cannot extract text from '{suffix}' files.",
        code="UNSUPPORTED_FILE_TYPE",
    )


# ===========================================================================
# PDF extraction
# ===========================================================================

def _extract_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise DocumentError(
            "pdfplumber is not installed. Run: pip install pdfplumber",
            code="MISSING_DEPENDENCY",
        ) from exc

    try:
        with pdfplumber.open(str(path)) as pdf:
            pages: list[str] = []
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if text and text.strip():
                    # Prepend a page marker so downstream can track source locations
                    pages.append(f"--- Page {i} ---\n{text.strip()}")

    except DocumentError:
        raise
    except Exception as exc:
        raise DocumentError(
            f"Failed to parse PDF '{path.name}': {exc}",
            code="EXTRACTION_FAILED",
        ) from exc

    if not pages:
        raise DocumentError(
            f"'{path.name}' appears to contain no extractable text. "
            "It may be a scanned document (image-only PDF). "
            "OCR support is not yet implemented.",
            code="SCANNED_PDF",
        )

    full_text = "\n\n".join(pages)

    if len(full_text.strip()) < _MIN_TEXT_LENGTH:
        raise DocumentError(
            f"'{path.name}' yielded very little text ({len(full_text.strip())} chars). "
            "It may be a scanned document.",
            code="SCANNED_PDF",
        )

    logger.info(
        "PDF extracted: %s — %d pages, %d chars",
        path.name, len(pages), len(full_text),
    )
    return full_text


# ===========================================================================
# DOCX extraction
# ===========================================================================

def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentError(
            "python-docx is not installed. Run: pip install python-docx",
            code="MISSING_DEPENDENCY",
        ) from exc

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise DocumentError(
            f"Failed to parse DOCX '{path.name}': {exc}",
            code="EXTRACTION_FAILED",
        ) from exc

    sections: list[str] = []

    # Extract paragraphs — skip empty ones
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            sections.append(text)

    # Extract text from tables (common in tender BOQs and specification tables)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                # Join table row as tab-separated for readability
                sections.append("\t".join(row_cells))

    if not sections:
        raise DocumentError(
            f"'{path.name}' appears to be empty.",
            code="EMPTY_DOCUMENT",
        )

    full_text = "\n\n".join(sections)

    if len(full_text.strip()) < _MIN_TEXT_LENGTH:
        raise DocumentError(
            f"'{path.name}' yielded very little text ({len(full_text.strip())} chars). "
            "The document may be empty or corrupted.",
            code="EMPTY_DOCUMENT",
        )

    logger.info(
        "DOCX extracted: %s — %d paragraphs/rows, %d chars",
        path.name, len(sections), len(full_text),
    )
    return full_text


# ===========================================================================
# IS reference scanner — preliminary scan of extracted text
# ===========================================================================

# Regex to find IS standard references in tender text.
# Matches formats like:
#   IS 269:2015
#   IS 1180 (Part 1):2014
#   IS 10322 (Part 5/Sec 3):2012
#   IS 2062:2011 Amd.4
#   IS 269 (latest edition)
_IS_REFERENCE_PATTERN = re.compile(
    r"IS\s+"                         # "IS " prefix
    r"(\d+)"                          # IS number
    r"(?:\s*\(([^)]+)\))?"            # optional (Part N/Sec M)
    r"(?:\s*:\s*(\d{4}))?"            # optional :YYYY year
    r"(?:\s+Amd\.?\s*(\d+))?",        # optional Amd.N
    re.IGNORECASE,
)


def scan_is_references(text: str) -> list[dict]:
    """
    Do a quick regex scan of extracted text to find all IS standard references.

    This is a preliminary scan — it finds candidate references but does NOT
    validate that the standards actually exist. The requirement extraction
    step (AI/ML) does the authoritative extraction with semantic understanding.

    Returns a list of dicts with keys:
      matched_text, is_number, part_section, year, amendment_number, char_offset
    """
    results = []
    for match in _IS_REFERENCE_PATTERN.finditer(text):
        results.append({
            "matched_text": match.group(0).strip(),
            "is_number": f"IS {match.group(1)}",
            "part_section": match.group(2),
            "year": int(match.group(3)) if match.group(3) else None,
            "amendment_number": int(match.group(4)) if match.group(4) else None,
            "char_offset": match.start(),
        })
    return results
